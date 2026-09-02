import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_left_border(cell, color_hex="4F46E5", size="36"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size) # 36 = ~4.5pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    for side in ['top', 'bottom', 'right']:
        n = OxmlElement(f'w:{side}')
        n.set(qn('w:val'), 'none')
        tcBorders.append(n)
        
    tcPr.append(tcBorders)

def build_document(output_path):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700
    
    # Custom Palette
    COLOR_PRIMARY = RGBColor(0x1E, 0x1B, 0x4B)   # Deep Indigo / Navy
    COLOR_ACCENT = RGBColor(0x4F, 0x46, 0xE5)    # Royal Indigo
    COLOR_SECONDARY = RGBColor(0x47, 0x55, 0x69) # Slate 600
    
    # Document Title / Cover Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("QuoteCompare Platform")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(20)
    run_sub = subtitle_p.add_run("Business Analyst Audit, Standard Operating Procedures (SOP), & Senior SaaS Architect Commercialization Blueprint")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = COLOR_ACCENT
    
    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document Ref:", "QC-SOP-ARCH-2026-V1"),
        ("Target Market:", "Australia (B2B & B2C Trade Services Marketplace)"),
        ("Role Authoring:", "Lead Business Analyst & Senior SaaS Platform Architect"),
        ("Date & Status:", "February 2026 | Approved Operational & Commercial Strategy")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(4.7)
        
        p0 = cell_lbl.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = COLOR_SECONDARY
        
        p1 = cell_val.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_PRIMARY
        
        set_cell_background(cell_lbl, 'F1F5F9')
        set_cell_background(cell_val, 'F8FAFC')
        set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_ACCENT
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_callout(text, title="KEY STRATEGIC TAKEAWAY"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "EEF2FF")
        add_left_border(cell, "4F46E5", "36")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(f"📌 {title}\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = COLOR_ACCENT
        
        r_b = p.add_run(text)
        r_b.font.size = Pt(10)
        r_b.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "0F172A") # Dark slate
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xF1, 0xF5, 0xF9)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # SECTION 1
    add_h1("1. Business Analyst (BA) Audit & Current System State")
    
    p = doc.add_paragraph("The QuoteCompare platform is a modern, web-based multi-role marketplace designed to bridge Australian buyers (homeowners, residential property managers, and businesses) with qualified trade vendors (electricians, solar/heat pump installers, air conditioning technicians, plumbers, and gas fitters).")
    p.paragraph_format.space_after = Pt(8)
    
    add_h2("1.1 Core Architecture & Technology Stack")
    p = doc.add_paragraph("The application leverages a modern micro-decoupled web architecture engineered for speed, scalability, and security:")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Frontend Client Layer: ").bold = True
    bp1.add_run("Angular 19+ (Standalone Components, Signals API for reactive state management, Vanilla CSS with custom glassmorphism and modern design tokens).")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Backend REST Service Layer: ").bold = True
    bp2.add_run("Node.js with Express framework, managing authentication verification, payload validation, custom routing, and transaction logic.")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("Database & Security Layer: ").bold = True
    bp3.add_run("Supabase PostgreSQL with granular Row Level Security (RLS) policies, native Auth JWT verification, and automated JSONB historical state tracking.")
    
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("Notification Dispatch: ").bold = True
    bp4.add_run("Nodemailer service providing automated HTML email confirmations, quote request receipt links, and status change alerts.")

    add_h2("1.2 Functional Capabilities Matrix")
    
    cap_table = doc.add_table(rows=6, cols=4)
    cap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Functional Module", "Target User Role", "Operational Status", "Key Capabilities Built"]
    
    hdr_row = cap_table.rows[0]
    for i, h_text in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, "1E1B4B")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    matrix_data = [
        ("Public Lead Capture", "Buyer / Guest", "Operational (V1.5.56)", "No-login required quote creation (/request-quote) with auto-generated 8-char short_id hash and tracking link."),
        ("Geo-Fenced Matching", "Vendor / Installer", "Operational (V1.5.56)", "Filtering of leads by State, Postcodes, Suburbs, and explicit Excluded Postcodes/Suburbs configured in vendor profile."),
        ("Interactive Comparison", "Buyer / Consumer", "Operational (V1.5.56)", "Side-by-side card grid showing vendor price, scope details, vendor identity, and complete price revision history tooltips."),
        ("Structured Negotiation", "Buyer & Vendor", "Operational (V1.5.56)", "Two-way negotiation loop (responded -> negotiating -> accepted) with automated JSONB history preservation."),
        ("Job Execution & Billing", "Vendor / Installer", "Operational (V1.5.56)", "Complete job workflow with mandatory invoice URL attachment for buyer delivery upon service fulfillment.")
    ]
    
    col_widths = [Inches(1.5), Inches(1.1), Inches(1.3), Inches(2.6)]
    
    for row_idx, data in enumerate(matrix_data, start=1):
        row = cap_table.rows[row_idx]
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 2
    add_h1("2. Standard Operating Procedures (SOP) for System Processes")
    
    p = doc.add_paragraph("This section defines standard operating workflows governing end-to-end interactions across Buyer, Vendor, and System Administrator personas.")
    
    add_h2("SOP 1: Buyer Quote Creation, Tracking & Account Auto-Claiming")
    add_h3("Objective: Allow buyers to easily post trade service requests and seamlessly track or claim leads.")
    
    p1 = doc.add_paragraph("1. Request Initiation: The buyer accesses /request-quote without requiring prior user authentication.")
    p2 = doc.add_paragraph("2. Specification Gathering: Buyer selects service category (e.g., Heat Pumps, Solar Panels, Battery Storage, Air Conditioning, Water Filters, EV Chargers, Electrical, Plumbing, Gas), inputs location details (Postcode, Suburb, State), and provides job details and PDF/image specifications.")
    p3 = doc.add_paragraph("3. Short ID Generation & Storage: The Node.js backend generates a cryptographic 8-character uppercase tracking hash (e.g., XK8W2M9N) and inserts the quote record into Supabase PostgreSQL.")
    p4 = doc.add_paragraph("4. Email & Tracking Link Dispatch: System sends a confirmation email containing direct link /track/:short_id.")
    p5 = doc.add_paragraph("5. Automated Claiming Mechanism: If an unauthenticated guest buyer subsequently signs up or logs in using the matching email address, the system automatically executes an ownership update, binding buyer_id = req.user.uid.")

    add_h2("SOP 2: Vendor Lead Discovery & Quoting Workflow")
    add_h3("Objective: Enable trade vendors to discover geo-targeted requests and submit transparent proposals.")
    
    p1 = doc.add_paragraph("1. Lead Discovery: Vendor logs into /vendor dashboard, accessing categorized request tabs (New Requests, Awaiting Response, Action Required, Completed).")
    p2 = doc.add_paragraph("2. Geo-Filtering Validation: System screens incoming open quotes against the vendor's profile settings (service_postcodes, service_suburbs, service_states) while suppressing leads matching excluded_postcodes.")
    p3 = doc.add_paragraph("3. Card Interaction: Clicking anywhere on a lead card opens the proposal submission form (/vendor/respond/:quoteId).")
    p4 = doc.add_paragraph("4. Quote Proposal Submission: Vendor inputs total price ($ AUD) and scope details. Submitting creates a record in quote_responses with status = 'responded'.")

    add_h2("SOP 3: Offer Comparison, Negotiation & Acceptance Lifecycle")
    add_h3("Objective: Provide a transparent negotiation loop for buyers and vendors.")
    
    p1 = doc.add_paragraph("1. Buyer Review: Buyer accesses /buyer/quote/:quoteId to inspect vendor offers side-by-side.")
    p2 = doc.add_paragraph("2. Renegotiation Request: If the buyer requests price or scope adjustment, they click 'Renegotiate', providing feedback. The system sets response status to 'negotiating' and writes a note in buyer_message.")
    p3 = doc.add_paragraph("3. Price History Archival: Upon vendor update, previous price, message, and timestamp are automatically appended to a JSONB history array, rendering an interactive price history tooltip on the frontend.")
    p4 = doc.add_paragraph("4. Final Acceptance: Buyer selects 'Accept', updating status to 'accepted'.")

    add_h2("SOP 4: Job Execution & Invoice Delivery Workflow")
    add_h3("Objective: Manage job completion and invoice link exchange.")
    
    p1 = doc.add_paragraph("1. Complete Job Trigger: Following physical work completion, vendor clicks 'Complete' under the 'Action Required' tab.")
    p2 = doc.add_paragraph("2. Invoice Verification: CompleteJobModal launches, requiring a valid invoice URL (e.g., Xero / MYOB hosted link or PDF).")
    p3 = doc.add_paragraph("3. State Transition: System updates quote_responses status to 'completed' and transfers the card to the 'Completed' dashboard tab.")

    add_callout(
        "All quote state transitions (Responded -> Negotiating -> Accepted -> Completed) maintain strict database validation and audit logs, ensuring zero data loss during re-negotiation.",
        "DATA INTEGRITY & AUDIT PROOFING"
    )

    # SECTION 3
    add_h1("3. Senior SaaS Platform Architect Blueprint – Australian Commercialization Strategy")
    
    p = doc.add_paragraph("To commercialize and scale the QuoteCompare platform across the Australian B2B and B2C trade services industry, the following technical, regulatory, tax, and architectural upgrades are recommended.")
    
    add_h2("3.1 Australian Regulatory, Tax & Compliance Framework")
    
    add_h3("A. Australian Business Number (ABN) Verification Engine")
    p = doc.add_paragraph("Unverified trade contractors create massive legal and financial liability under Australian Consumer Law (ACL). The platform must integrate directly with the Australian Business Register (ABR) Web Services API.")
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Workflow: ").bold = True
    bp1.add_run("During vendor registration, vendor inputs ABN -> Backend queries ABR API -> Verifies Entity Status (Active), GST Registration, and Business Name -> Grants 'ABN Verified' badge.")

    add_h3("B. Goods and Services Tax (GST 10%) Quote Breakdown")
    p = doc.add_paragraph("Under ATO guidelines and ACCC transparent pricing rules, quotes provided to Australian consumers must explicitly itemize 10% GST or state GST-inclusive totals.")

    add_h3("C. Privacy Act 1988 & APPs Compliance")
    p = doc.add_paragraph("Ensure all customer lead data is hosted in the ap-southeast-2 (Sydney) AWS region. Implement explicit opt-in consent for lead sharing under Australian Privacy Principles (APP 3 & APP 6).")

    add_h2("3.2 Commercial Monetization & Revenue Models")
    
    mon_table = doc.add_table(rows=4, cols=3)
    mon_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_headers = ["Revenue Stream", "Pricing Architecture", "Commercial Implementation Strategy"]
    
    m_hdr_row = mon_table.rows[0]
    for i, h_text in enumerate(m_headers):
        cell = m_hdr_row.cells[i]
        set_cell_background(cell, "1E1B4B")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    m_data = [
        ("Pay-Per-Lead (PPL) Unlocks", "$25.00 - $45.00 AUD per lead unlock", "Vendors purchase credit packs. Unlocking buyer contact information (phone/email) deducts credits based on trade category complexity (e.g. Solar = 3 credits, Plumbing = 1 credit)."),
        ("Tiered Vendor Subscriptions", "Starter ($49/mo) | Growth ($149/mo) | Pro ($399/mo AUD)", "Stripe Subscription Billing integration. Higher tiers grant larger postcode radius coverage, instant SMS alerts, and priority placement in quote buyer views."),
        ("Escrow & Milestone Processing", "2.5% to 5.0% platform take-rate", "Stripe Connect escrow holds buyer milestone deposits compliant with State Security of Payments Acts (NSW SOPA / VIC SOPA), releasing payments upon buyer job sign-off.")
    ]
    
    m_widths = [Inches(1.8), Inches(1.8), Inches(2.9)]
    for row_idx, data in enumerate(m_data, start=1):
        row = mon_table.rows[row_idx]
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = m_widths[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_h2("3.3 Enterprise Technical Roadmap")
    
    add_h3("Phase 1: Real-Time Communication & SMS Notification Engine")
    p = doc.add_paragraph("Email notifications suffer from low open rates among active field contractors. Integrating WebSockets (Supabase Realtime) and an SMS gateway (Twilio / MessageMedia) will enable instant lead dispatches via SMS.")

    add_h3("Phase 2: Xero & MYOB OAuth 2.0 Accounting Sync")
    p = doc.add_paragraph("Enable two-way synchronization with Xero and MYOB. When a buyer accepts a quote, the platform pushes a Draft Invoice directly into the vendor's accounting system.")

    add_h3("Phase 3: AI Document Vision & Job Spec Parsing")
    p = doc.add_paragraph("Implement AI Document Vision (Claude / Vision AI) to analyze buyer-uploaded photographs (e.g. electrical switchboards, roof layouts) and automatically extract specs to populate lead requirements.")

    add_h2("3.4 Database Schema Migration Script (Production Commercialization)")
    p = doc.add_paragraph("Below is the recommended PostgreSQL DDL script to support ABN verification, GST itemization, credit tracking, and Stripe commercial integration:")

    sql_code = """-- Commercialization Schema Migration Script
ALTER TABLE public.profiles
ADD COLUMN IF NOT EXISTS abn_verified boolean DEFAULT false,
ADD COLUMN IF NOT EXISTS abn_business_name text,
ADD COLUMN IF NOT EXISTS stripe_customer_id text,
ADD COLUMN IF NOT EXISTS stripe_subscription_status text DEFAULT 'inactive',
ADD COLUMN IF NOT EXISTS available_lead_credits integer DEFAULT 0;

ALTER TABLE public.quotes
ADD COLUMN IF NOT EXISTS preferred_start_date date,
ADD COLUMN IF NOT EXISTS site_inspection_required boolean DEFAULT false,
ADD COLUMN IF NOT EXISTS property_ownership text CHECK (property_ownership IN ('owner', 'renter', 'commercial'));

ALTER TABLE public.quote_responses
ADD COLUMN IF NOT EXISTS subtotal numeric(10,2),
ADD COLUMN IF NOT EXISTS gst_amount numeric(10,2),
ADD COLUMN IF NOT EXISTS est_completion_days integer,
ADD COLUMN IF NOT EXISTS warranty_terms text,
ADD COLUMN IF NOT EXISTS deposit_required_percent numeric(5,2) DEFAULT 0.00;"""

    add_code_block(sql_code)

    # Save document
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    import sys
    out = sys.argv[1] if len(sys.argv) > 1 else "QuoteCompare_Platform_SOP_and_Architectural_Blueprint.docx"
    build_document(out)
